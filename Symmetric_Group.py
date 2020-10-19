# 2020-07-29
from docx import Document
import random


def quest_int(question, quest_min, quest_max, quest_obj):
    try:
        quest_val = int(input(f'{question}'))
        if quest_min != 0 and quest_val < quest_min:
            print(f'Need to use at least {quest_min} {quest_obj}')
            raise ValueError
        if quest_max != 0 and quest_max < quest_val:
            print(f"Can't use more than {quest_max} {quest_obj}")
            raise ValueError
    except ValueError:
        quest_val = quest_int(question, quest_min, quest_max, quest_obj)
    return quest_val


def list_to_num(list_to_update):
    number = 0
    for value in list_to_update:
        number = number * 10 + value
    return number


def cycle_single_creator(cycle_length):
    cycle = []
    numbers_choice = list(range(1, sym_size + 1))
    while len(cycle) < cycle_length:
        random_element = random.choice(numbers_choice)
        numbers_choice.remove(random_element)
        cycle.append(random_element)

    return cycle


def cycle_multiple_creator():
    global question_attempt_type
    cycle_dict_tally = {}  # Keeps track of number of times each cycle occurs
    question_index_value = []  # Location of each cycle in cycle_dict_tally
    total_cycles = random.randint(2, cycle_tot)
    cycles_list = []
    cycle_count = 0
    cycle_attempts = 0
    while cycle_count < total_cycles and cycle_attempts < total_attempts:
        cycle_length = random.randint(2, cycle_maxlen)
        cycle = cycle_single_creator(cycle_length)

        cycle_index_value = list_to_num(cycle)
        if cycle_index_value in cycle_dict_tally:
            if cycle_dict_tally[cycle_index_value] == cycle_length - 1:
                cycle_attempts += 1
                continue
            cycle_dict_tally[cycle_index_value] = cycle_dict_tally[cycle_index_value] + 1
        else:
            cycle_dict_tally[cycle_index_value] = 1
            if cycle_index_value not in cycles_full_index:
                cycles_full_index.append(cycle_index_value)

        cycles_list.append(cycle)

        cycle_count += 1

        cycle_index_location = cycles_full_index.index(cycle_index_value)
        question_index_value.append(cycle_index_location)

    if cycles_disjoint_check(cycles_list) == 1:  # Already solved
        question_attempt_type = 1
        return cycles_list

    # q_index_value is a dummy value
    if all(q_index_value != question_index_value for q_index_value in questions_full_index):
        questions_full_index.append(question_index_value)
    elif question_attempts != total_attempts:
        question_attempt_type = 2

    return cycles_list


def cycle_element_mapping(element_input, cycle_function):
    if element_input in cycle_function:
        element_index_loc = cycle_function.index(element_input)
        if element_index_loc == len(cycle_function) - 1:
            element_input = cycle_function[0]
        else:
            element_input = cycle_function[element_index_loc + 1]
    # else we leave the element as it is
    return element_input  # Now the output


def answer_question(cycles_list):
    cycles_list = cycles_list[::-1]
    # NOTE: Composition is now done from left to right.
    numbers_choice = list(range(1, sym_size + 1))
    elements_found = []
    cycles_solution = []
    for element_count, element_value in enumerate(numbers_choice):
        cycle_sub_solution = []
        while element_value not in elements_found:
            cycle_sub_solution.append(element_value)
            elements_found.append(element_value)
            for cycle_sub in cycles_list:  # finds where element eventually maps to
                element_value = cycle_element_mapping(element_value, cycle_sub)
            if element_value == cycle_sub_solution[0]:
                if len(cycle_sub_solution) != 1:
                    cycles_solution.append(cycle_sub_solution)
                break
    if len(cycles_solution) == 0:
        cycles_solution.append([])
    return cycles_solution


def cycles_disjoint_check(cycles_list):
    disjoint_test = 1
    elements_found = []
    for cycle in cycles_list:
        for element in cycle:
            if element in elements_found:
                disjoint_test = 0
                return disjoint_test
            elements_found.append(element)
    return disjoint_test


def write_question(question, answer):
    for result in [question, answer]:

        if result == question:
            current_file = Question_file
        else:
            current_file = Answer_file

        if not result[0]:
            current_file.add_paragraph(f"{quest_count}:     1")
            continue
        else:
            p = current_file.add_paragraph(f"{quest_count}:    (")

        for cycle_count, cycle in enumerate(result):
            if cycle_count != 0:
                p.add_run('(')
            for element in cycle:
                p.add_run(f"{element}")
                if element != cycle[-1]:
                    p.add_run(' ')
            p.add_run(')')


sym_size = quest_int(
    'How many elements do you want to work with?: ', quest_min=3, quest_max=9, quest_obj='elements'
)
cycle_maxlen = quest_int(
    'What is the maximum length of any given cycle?: ', quest_min=2, quest_max=sym_size, quest_obj='elements'
)
cycle_tot = quest_int(
    'What''s the maximum number of cycles you want to work with?: ', quest_min=2, quest_max=0, quest_obj='cycles'
)
quest_tot = quest_int(
    'How many questions would you like to have?: ', quest_min=1, quest_max=0, quest_obj='questions'
)

Question_file = Document()
Question_file.add_heading('Symmetric Group Composition', 0)
Question_file.add_heading('Questions', level=1)

Answer_file = Document()
Answer_file.add_heading('Symmetric Group Composition', 0)
Answer_file.add_heading('Answers', level=1)

cycles_full_index = []  # Holds permutations of cycles found during all creations
questions_full_index = []  # Uses indices of cycles_full_index values as pointers for order
quest_count = 1
question_attempts = 0
question_attempt_type = 0  # 1 if question already solved, 2 if it's previous question
total_attempts = 10
while quest_count <= quest_tot:
    question_cycles = cycle_multiple_creator()
    if question_attempt_type != 0:  # Isn't a new question we can use
        if question_attempt_type == 2:  # We have a previous asked question
            question_attempts += 1
    if question_attempts == total_attempts:
        break
    if question_attempt_type != 0:
        question_attempt_type = 0
        continue

    answer_cycles = answer_question(question_cycles)

    write_question(question_cycles, answer_cycles)

    # print(*question_cycles)
    # print(*answer_cycles)
    # print('--------------------------')
    quest_count += 1

# print(cycles_full_index)

Question_file.save('Symmetric_Questions.docx')
Answer_file.save('Symmetric_Answers.docx')
